import tkinter as tk
from tkinter import ttk
import tkinter.messagebox
import pickle
import pandas as pd
import matplotlib
import calendar
import matplotlib.pyplot as plt
import seaborn as sns
from PIL import Image
from wordcloud import WordCloud
from docx import Document
from docx.shared import Inches
matplotlib.use('TkAgg')
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
#first window
window = tk.Tk()
window.title('Welcome to Corona Virus 2019 information platform')
window.geometry('800x500')
pd.set_option('display.max_columns', None)
#set background
canvas = tk.Canvas(window, height=1000, width=800)
imagefile = tk.PhotoImage(file='main.png')
image = canvas.create_image(0, 0, anchor='nw', image=imagefile)
canvas.pack(side='top')
tk.Label(window,text = 'Welcome to our system! Click User know more about COVID-19').place(x=40,y=30)

#get the information form the dataset
data = pd.read_csv('covid_19_data.csv')
country1 = data['Country/Region'].drop_duplicates()
pd.set_option('display.max_rows', None)

country1 = country1.tolist()
country = data.drop_duplicates(['Country/Region', 'Province'])
province = country.groupby('Country/Region').Province.apply(list).to_dict()

# define function for user interface

def userframe():
    window_user = tk.Toplevel(window)
    window_user.geometry('300x300')
    window_user.title('Enter details you want to enquire')
    pro = tk.StringVar()
    # when the country is selected, show the province
    def go(*args):
        tk.Label(window_user, text='Province').place(x=40, y=180)

        combox_province = ttk.Combobox(window_user, textvariable=pro)
        combox_province['values'] = province[cy.get()]
        combox_province.place(x=120, y=180)
        tk.Label(window_user, text='then information would be about country').place(x=40, y=150)

    # information about country
    L1 = tk.Label(window_user, text='Country')
    L1.place(x=40, y=100)
    cy = tk.StringVar()
    combox_country = ttk.Combobox(window_user, textvariable=cy)
    combox_country['values'] = country1
    combox_country.current(0)
    combox_country.bind("<<ComboboxSelected>>", go)
    combox_country.place(x=120, y=100)
    world = tk.StringVar()
    tk.Label(window_user,text ="Click CheckButton Analysis The World ! ").place(x=10,y=10)
    tk.Label(window_user,text="Otherwise Analysis By Country!").place(x=10,y=40)
    testworld = tk.StringVar()
    def change():

        if testworld.get()=='1':
            combox_country.place_forget()
            L1.place_forget()
        else:
            combox_country.place(x=120, y=100)
            L1.place(x=40, y=100)

    c = tk.Checkbutton(window_user,text = "Analysis By The world",variable =testworld,command=change)
    c.place(x=10,y=60)




    def Analysis():
        print(testworld.get())
        if testworld.get() == '1':
            print("Analysis by world")
            window_user_info = tk.Toplevel(window_user)
            window_user_info.geometry('1300x1000')

            if pro.get() == '':

                # no province input, then anaalys the country

                #total_country = data[data['Country/Region'] == cy.get()]
                total_country=data

                # total_confirm= total_confirm['Confirmed'].tail(0)
                total_confirm = total_country.groupby(['ObservationDate'], sort=False)['Confirmed'].sum()

                time1 = total_confirm.index
                num_confirm = int(total_confirm.tail(1))
                date = str(total_confirm.tail(1).index[0])

                # get the number of Death
                total_death = total_country.groupby(['ObservationDate'], sort=False)['Deaths'].sum()
                time2 = total_death.index
                num_death = int(total_death.tail(1))
                # get the number of Recovered
                total_recover = total_country.groupby(['ObservationDate'], sort=False)['Recovered'].sum()
                time3 = total_recover.index
                num_recover = int(total_recover.tail(1))
                # plot for confired  recovered death
                f = Figure(figsize=(3.5, 2.3), dpi=100)
                f_plot = f.add_subplot(111)
                canvs = FigureCanvasTkAgg(f, window_user_info)
                canvs.get_tk_widget().place(x=50, y=100)
                f_plot.plot(time1, total_confirm, label="Confirmed ")
                f_plot.plot(time2, total_death, label="Death")
                f_plot.plot(time3, total_recover, label="Recover")
                f_plot.legend()
                f.savefig('1')
                canvs.draw()
                # show data information
                pie = plt.figure(figsize=(3, 3), facecolor="#F0F0F0")
                #
                # get data

                region = data
                region = region[region['ObservationDate'] == date]
                #print(region)
                region = region.groupby(['Country/Region'])['Confirmed'].sum()
               # print(type(region))
                region =region.sort_values()
                countries = region.index

                position = []
                data_confrimed = []

                for  row in region:

                    data_confrimed.append(row)
                for  row in countries:
                    position.append(row)

                pie.patches, pie.text2, pie.text1 = plt.pie(data_confrimed,
                                                            labels=position,
                                                            autopct='%3.1f%%',
                                                            startangle=90,
                                                            pctdistance=1.4,
                                                            textprops={'fontsize': 8, 'color': '#000080'}
                                                            )
                plt.axis('equal')
                plt.savefig('2')

                canvas_statis = FigureCanvasTkAgg(pie, window_user_info)
                canvas_statis.get_tk_widget().place(x=890, y=60)
                # for confirmed
                tk.Label(window_user_info, text="The Whole World", font=("'Arial", 15), fg="black").place(x=30, y=10)
                tk.Label(window_user_info, text=date, font=("'Arial", 15), fg="black").place(x=30, y=50)
                tk.Label(window_user_info, text='Confirmed:', font=("'Arial", 15), fg="red").place(x=160, y=50)
                tk.Label(window_user_info, text=num_confirm, font=("'Arial", 15), fg="red").place(x=260, y=50)
                # for death
                tk.Label(window_user_info, text='Death:', font=("'Arial", 15), fg="#FF6600").place(x=420, y=50)
                tk.Label(window_user_info, text=num_death, font=("'Arial", 15), fg="#FF6600").place(x=480, y=50)
                # for recover
                tk.Label(window_user_info, text='Recovered:', font=("'Arial", 15), fg="#FF6600").place(x=640, y=50)
                tk.Label(window_user_info, text=num_recover, font=("'Arial", 15), fg="#FF6600").place(x=740, y=50)
                # form the data from the data set

                data_country = []
                # print(region)
                region = data
                region = region[region['ObservationDate'] == date]
                # print(region)
                region = region.groupby(['Country/Region'])['Confirmed','Deaths','Recovered'].sum()
                region = region.sort_values('Confirmed')
                print(region)
                dataTreeView = ttk.Treeview(window_user_info, show='headings', column=('RG', 'FR', 'CF', 'DT', 'RC'))
                dataTreeView.column('RG', width=80, anchor="center")
                dataTreeView.column('FR', width=80, anchor="center")
                dataTreeView.column('CF', width=80, anchor="center")
                dataTreeView.column('DT', width=80, anchor="center")
                dataTreeView.column('RC', width=80, anchor="center")
                dataTreeView.heading('RG', text='Region')
                dataTreeView.heading('FR', text='Fatality rate')
                dataTreeView.heading('CF', text='Confirmed')
                dataTreeView.heading('DT', text='Death')
                dataTreeView.heading('RC', text='Recover')
                dataTreeView.place(x=440, y=100)
                for index in region.index:
                    one = []
                    one.append(index)
                    FR = format((region.loc[index].values[1] / region.loc[index].values[2]) * 100, '.3f')
                    FR = str(FR) + '%'
                    one.append(FR)
                    one.append(region.loc[index].values[0])
                    one.append(region.loc[index].values[1])
                    one.append(region.loc[index].values[2])
                    dataTreeView.insert("", 1, text="line1", values=one)
                # for general infromation

                tk.Label(window_user_info, text='General Information:', font=("'Arial", 15)).place(x=30, y=360)
                # Get information about age
                data1 = pd.read_csv('COVID19_line_list_data.csv')
                age = data1['age']
                # draw the age diagram
                age = age.dropna()
                f_age = Figure(figsize=(3.5, 2.3), dpi=100)
                f_plot_age = f_age.add_subplot(111)
                canvs2 = FigureCanvasTkAgg(f_age, window_user_info)
                canvs2.get_tk_widget().place(x=50, y=400)
                f_plot_age.hist(age)
                f_plot_age.set_title('Age of illness')
                f_age.savefig('3')
                canvs2.draw()
                tk.Label(window_user_info, text='Min Age:', font=("'Arial", 10)).place(x=30, y=640)
                tk.Label(window_user_info, text=age.min(), font=("'Arial", 10)).place(x=90, y=640)
                tk.Label(window_user_info, text='Max Age:', font=("'Arial", 10)).place(x=140, y=640)
                tk.Label(window_user_info, text=age.max(), font=("'Arial", 10)).place(x=200, y=640)
                tk.Label(window_user_info, text='Mean Age:', font=("'Arial", 10)).place(x=260, y=640)
                tk.Label(window_user_info, text=format(age.mean(), '.1f'), font=("'Arial", 10)).place(x=320, y=640)

                # get infromation about wuhan;
                wuhan_visit = data1['visiting Wuhan'].dropna()
                num_visit = 0
                num_novisit = 0
                for people in wuhan_visit:
                    # print(people)
                    if people == 1:
                        num_visit += 1
                    elif people == 0:
                        num_novisit += 1
                wuhan = [num_visit, num_novisit]
                go_wuhan = ['Went Wuhan', 'Not went Wuhan']
                f_wuhan = Figure(figsize=(3.5, 2.3), dpi=100)
                f_plot_wuhan = f_wuhan.add_subplot(111)
                canvs3 = FigureCanvasTkAgg(f_wuhan, window_user_info)
                canvs3.get_tk_widget().place(x=460, y=400)
                f_plot_wuhan.pie(wuhan, labels=go_wuhan, autopct='%3.1f%%',
                                 shadow=True,
                                 startangle=90,
                                 pctdistance=1.4,
                                 textprops={'fontsize': 8, 'color': '#000080'})
                f_plot_wuhan.set_title('Percent of People been to Wuhan')
                f_wuhan.savefig('4')
                canvs3.draw()
                # find the world could for symptom
                data2 = pd.read_csv('COVID19_open_line_list.csv')
                syp = data2['symptoms'].dropna()
                description = " ".join(symptoms for symptoms in syp)
                worldcould = WordCloud().generate(description)
                fig = plt.figure(figsize=(3.5, 2.3))
                im = plt.imshow(worldcould, interpolation='bilinear')
                plt.title('WorldCould')
                plt.axis("off")
                plt.savefig('5')
                canvs4 = FigureCanvasTkAgg(fig, window_user_info)
                canvs4.get_tk_widget().place(x=860, y=400)
                canvs4.draw()

                def printprot():
                    doc = Document()
                    doc.add_paragraph('The report is about :' + cy.get())
                    doc.add_paragraph('The total number of Confirmed people: ' + str(num_confirm))
                    doc.add_paragraph('The total number of Death people: ' + str(num_death))
                    doc.add_paragraph('The total number of Recover people: ' + str(num_death))
                    doc.add_picture('1.png', width=Inches(2))
                    doc.add_paragraph('Here is the situation about region: ')
                    doc.add_paragraph(region.to_string())
                    doc.add_picture('2.png', width=Inches(2))
                    doc.add_paragraph('The average age of illness is ' + str(age.mean()))
                    doc.add_paragraph('The maxmium age of illness is ' + str(age.max()))
                    doc.add_paragraph('the minmium age of illness is ' + str(age.min()))
                    doc.add_picture('3.png')
                    doc.add_paragraph('Here is the picture about percentage of people who went to Wuhan ')
                    doc.add_paragraph('4.png')
                    doc.add_paragraph('Here is the picture of worldCould')
                    doc.add_picture('5.png')

                    doc.save('word.docx')

                    tk.messagebox.showinfo('Successful', 'see the report!')

                print_report = tk.Button(window_user_info, text='Print Report', command=printprot, bg='#FF9900',
                                         height=2, width=8,
                                         relief='raised')
                print_report.place(x=600, y=630)







            else:
                print('province we have')

        else:
            print("Analysis by coutnry")
            window_user_info = tk.Toplevel(window_user)
            window_user_info.geometry('1300x1000')

            if pro.get() == '':

                # no province input, then anaalys the country

                total_country = data[data['Country/Region'] == cy.get()]
                # total_confirm= total_confirm['Confirmed'].tail(0)
                total_confirm = total_country.groupby(['ObservationDate'], sort=False)['Confirmed'].sum()
                time1 = total_confirm.index
                num_confirm = int(total_confirm.tail(1))
                date = str(total_confirm.tail(1).index[0])

                # get the number of Death
                total_death = total_country.groupby(['ObservationDate'], sort=False)['Deaths'].sum()
                time2 = total_death.index
                num_death = int(total_death.tail(1))
                # get the number of Recovered
                total_recover = total_country.groupby(['ObservationDate'], sort=False)['Recovered'].sum()
                time3 = total_recover.index
                num_recover = int(total_recover.tail(1))
                # plot for confired  recovered death
                f = Figure(figsize=(3.5, 2.3), dpi=100)
                f_plot = f.add_subplot(111)
                canvs = FigureCanvasTkAgg(f, window_user_info)
                canvs.get_tk_widget().place(x=50, y=100)
                f_plot.plot(time1, total_confirm, label="Confirmed ")
                f_plot.plot(time2, total_death, label="Death")
                f_plot.plot(time3, total_recover, label="Recover")
                f_plot.legend()
                f.savefig('1')

                canvs.draw()
                # show data information
                pie = plt.figure(figsize=(3, 3), facecolor="#F0F0F0")
                #
                # get data
                region = total_country
                region = region[region['ObservationDate'] == date]
                region = region.sort_values('Confirmed')

                position = []
                data_confrimed = []
                for index, row in region.iterrows():
                    position.append(row['Province'])
                    data_confrimed.append(row['Confirmed'])

                pie.patches, pie.text2, pie.text1 = plt.pie(data_confrimed,
                                                            labels=position,
                                                            autopct='%3.1f%%',
                                                            shadow=True,
                                                            startangle=90,
                                                            pctdistance=1.4,
                                                            textprops={'fontsize': 8, 'color': '#000080'}
                                                            )
                plt.axis('equal')
                plt.savefig('2')

                canvas_statis = FigureCanvasTkAgg(pie, window_user_info)
                canvas_statis.get_tk_widget().place(x=890, y=60)
                # for confirmed
                tk.Label(window_user_info, text=cy.get(), font=("'Arial", 15), fg="black").place(x=30, y=10)
                tk.Label(window_user_info, text=date, font=("'Arial", 15), fg="black").place(x=30, y=50)
                tk.Label(window_user_info, text='Confirmed:', font=("'Arial", 15), fg="red").place(x=160, y=50)
                tk.Label(window_user_info, text=num_confirm, font=("'Arial", 15), fg="red").place(x=260, y=50)
                # for death
                tk.Label(window_user_info, text='Death:', font=("'Arial", 15), fg="#FF6600").place(x=420, y=50)
                tk.Label(window_user_info, text=num_death, font=("'Arial", 15), fg="#FF6600").place(x=480, y=50)
                # for recover
                tk.Label(window_user_info, text='Recovered:', font=("'Arial", 15), fg="#FF6600").place(x=640, y=50)
                tk.Label(window_user_info, text=num_recover, font=("'Arial", 15), fg="#FF6600").place(x=740, y=50)
                # form the data from the data set

                data_country = []
                # print(region)
                dataTreeView = ttk.Treeview(window_user_info, show='headings', column=('RG', 'FR', 'CF', 'DT', 'RC'))
                dataTreeView.column('RG', width=80, anchor="center")
                dataTreeView.column('FR', width=80, anchor="center")
                dataTreeView.column('CF', width=80, anchor="center")
                dataTreeView.column('DT', width=80, anchor="center")
                dataTreeView.column('RC', width=80, anchor="center")
                dataTreeView.heading('RG', text='Region')
                dataTreeView.heading('FR', text='Fatality rate')
                dataTreeView.heading('CF', text='Confirmed')
                dataTreeView.heading('DT', text='Death')
                dataTreeView.heading('RC', text='Recover')
                dataTreeView.place(x=440, y=100)
                for index in region.index:
                    one = []
                    one.append(region.loc[index].values[2])
                    FR = format((region.loc[index].values[6] / region.loc[index].values[5]) * 100, '.3f')
                    FR = str(FR) + '%'
                    one.append(FR)
                    one.append(region.loc[index].values[5])
                    one.append(region.loc[index].values[6])
                    one.append(region.loc[index].values[7])
                    dataTreeView.insert("", 1, text="line1", values=one)
                # for general infromation

                tk.Label(window_user_info, text='General Information:', font=("'Arial", 15)).place(x=30, y=360)
                # Get information about age
                data1 = pd.read_csv('COVID19_line_list_data.csv')
                age = data1['age']
                # draw the age diagram
                age = age.dropna()
                f_age = Figure(figsize=(3.5, 2.3), dpi=100)
                f_plot_age = f_age.add_subplot(111)
                canvs2 = FigureCanvasTkAgg(f_age, window_user_info)
                canvs2.get_tk_widget().place(x=50, y=400)
                f_plot_age.hist(age)
                f_plot_age.set_title('Age of illness')
                f_age.savefig('3')
                canvs2.draw()
                tk.Label(window_user_info, text='Min Age:', font=("'Arial", 10)).place(x=30, y=640)
                tk.Label(window_user_info, text=age.min(), font=("'Arial", 10)).place(x=90, y=640)
                tk.Label(window_user_info, text='Max Age:', font=("'Arial", 10)).place(x=140, y=640)
                tk.Label(window_user_info, text=age.max(), font=("'Arial", 10)).place(x=200, y=640)
                tk.Label(window_user_info, text='Mean Age:', font=("'Arial", 10)).place(x=260, y=640)
                tk.Label(window_user_info, text=format(age.mean(), '.1f'), font=("'Arial", 10)).place(x=320, y=640)

                # get infromation about wuhan;
                wuhan_visit = data1['visiting Wuhan'].dropna()
                num_visit = 0
                num_novisit = 0
                for people in wuhan_visit:
                    # print(people)
                    if people == 1:
                        num_visit += 1
                    elif people == 0:
                        num_novisit += 1
                wuhan = [num_visit, num_novisit]
                go_wuhan = ['Went Wuhan', 'Not went Wuhan']
                f_wuhan = Figure(figsize=(3.5, 2.3), dpi=100)
                f_plot_wuhan = f_wuhan.add_subplot(111)
                canvs3 = FigureCanvasTkAgg(f_wuhan, window_user_info)
                canvs3.get_tk_widget().place(x=460, y=400)
                f_plot_wuhan.pie(wuhan, labels=go_wuhan, autopct='%3.1f%%',
                                 shadow=True,
                                 startangle=90,
                                 pctdistance=1.4,
                                 textprops={'fontsize': 8, 'color': '#000080'})
                f_plot_wuhan.set_title('Percent of People been to Wuhan')
                f_wuhan.savefig('4')
                canvs3.draw()
                # find the world could for symptom
                data2 = pd.read_csv('COVID19_open_line_list.csv')
                syp = data2['symptoms'].dropna()
                description = " ".join(symptoms for symptoms in syp)
                worldcould = WordCloud().generate(description)
                fig = plt.figure(figsize=(3.5, 2.3))
                im = plt.imshow(worldcould, interpolation='bilinear')
                plt.title('WorldCould')
                plt.axis("off")
                plt.savefig('5')
                canvs4 = FigureCanvasTkAgg(fig, window_user_info)
                canvs4.get_tk_widget().place(x=860, y=400)
                canvs4.draw()

                def printprot():
                    doc = Document()
                    doc.add_paragraph('The report is about :' + cy.get())
                    doc.add_paragraph('The total number of Confirmed people: ' + str(num_confirm))
                    doc.add_paragraph('The total number of Death people: ' + str(num_death))
                    doc.add_paragraph('The total number of Recover people: ' + str(num_death))
                    doc.add_picture('1.png', width=Inches(2))
                    doc.add_paragraph('Here is the situation about region: ')
                    doc.add_paragraph(region.to_string())
                    doc.add_picture('2.png', width=Inches(2))
                    doc.add_paragraph('The average age of illness is ' + str(age.mean()))
                    doc.add_paragraph('The maxmium age of illness is ' + str(age.max()))
                    doc.add_paragraph('the minmium age of illness is ' + str(age.min()))
                    doc.add_picture('3.png')
                    doc.add_paragraph('Here is the picture about percentage of people who went to Wuhan ')
                    doc.add_paragraph('4.png')
                    doc.add_paragraph('Here is the picture of worldCould')
                    doc.add_picture('5.png')

                    doc.save('word.docx')

                    tk.messagebox.showinfo('Successful', 'see the report!')

                print_report = tk.Button(window_user_info, text='Print Report', command=printprot, bg='#FF9900',
                                         height=2, width=8,
                                         relief='raised')
                print_report.place(x=600, y=630)







            else:
                print('province we have')

    tk.Button(window_user,command = Analysis, text='Analysis', bg='#FF9900', height=2, width=8, relief='raised').place(x=180,y=200)


bt_user = tk.Button(window,text='User',command = userframe,bg='#FF9900',height = 2,width =8,relief ='raised' )
bt_user.place(x=40,y=60)
def admin_sign():
    window_admin = tk.Toplevel(window)
    window_admin.geometry('350x200')
    window_admin.title('Sign up as Admin')
    new_name = tk.StringVar()
    tk.Label(window_admin, text='Admin:').place(x=60, y=40)
    tk.Label(window_admin, text='Password:').place(x=60, y=80)
    # Admin input
    var_usr_name = tk.StringVar()
    entry_usr_name = tk.Entry(window_admin, textvariable=var_usr_name)
    entry_usr_name.place(x=160, y=40)
    # password input
    var_usr_pwd = tk.StringVar()
    entry_usr_pwd = tk.Entry(window_admin, textvariable=var_usr_pwd, show='*')
    entry_usr_pwd.place(x=160, y=80)

    def usr_log_in():
        # get user name and user from input
        usr_name = var_usr_name.get()
        usr_pwd = var_usr_pwd.get()
        # check whether we have the database if not we would create one
        try:
            with open('usr_info.pickle', 'rb') as usr_file:
                usrs_info = pickle.load(usr_file)
        except FileNotFoundError:
            with open('usr_info.pickle', 'wb') as usr_file:
                usrs_info = {'admin': 'admin'}
                pickle.dump(usrs_info, usr_file)
        # check whether we have user information and and passwords
        if usr_name in usrs_info:
            if usr_pwd == usrs_info[usr_name]:
                tk.messagebox.showinfo(title='welcome',
                                       message='welcome：' + usr_name)
                #go to the admin page
                window_admin_change = tk.Toplevel(window_admin)
                window_admin_change.geometry('300x200')
                #test
                window_admin_change.title('Change country')
                change_country = tk.StringVar()
                tk.Label(window_admin_change, text='Country:').place(x=80, y=10)
                tk.Entry(window_admin_change, textvariable=change_country).place(x=150, y=10)
                def delete_country():
                    if change_country.get() in country1:
                        country1.remove(change_country.get())
                        tk.messagebox.showinfo('Successful', 'deleted')
                    else:
                        tk.messagebox.showerror('Error', 'Country is not existed')
                tk.Button(window_admin_change, text='delete',command=delete_country).place(x=180,y=100)
                def add_country():
                    country1.append(change_country.get())
                    tk.messagebox.showinfo('Successful', 'added')
                tk.Button(window_admin_change,text='add',command=add_country).place(x=120,y=100)


            else:
                tk.messagebox.showerror(message='password is not correct')
        # password username can not be empty
        elif usr_name == '' or usr_pwd == '':
            tk.messagebox.showerror(message='user and password can not be empty')
        # if it is not in database, check in the database
        else:
            is_signup = tk.messagebox.askyesno('welcome', 'you are not registered yet! register?')
            if is_signup:
                usr_sign_up()

    def usr_sign_up():
        # run this function
        def signtowcg():
            # get the information from input
            nn = new_name.get()
            np = new_pwd.get()
            npf = new_pwd_confirm.get()

            #get user information
            try:
                with open('usr_info.pickle', 'rb') as usr_file:
                    exist_usr_info = pickle.load(usr_file)
            except FileNotFoundError:
                exist_usr_info = {}

                # whether we have in the database and empty
            if nn in exist_usr_info:
                tk.messagebox.showerror('Error', 'User is already used')
            elif np == '' or nn == '':
                tk.messagebox.showerror('Error', 'Empty for user or password')
            elif np != npf:
                tk.messagebox.showerror('Error', 'Two different passwords')
            # if there is nothing wrong, put information into database
            else:
                exist_usr_info[nn] = np
                with open('usr_info.pickle', 'wb') as usr_file:
                    pickle.dump(exist_usr_info, usr_file)
                tk.messagebox.showinfo('welcome', 'successful')
                # close sign up frame
                window_sign_up.destroy()

        # register frame
        window_sign_up = tk.Toplevel(window)
        window_sign_up.geometry('350x200')
        window_sign_up.title('register')
        # username input
        new_name = tk.StringVar()
        tk.Label(window_sign_up, text='User：').place(x=10, y=10)
        tk.Entry(window_sign_up, textvariable=new_name).place(x=150, y=10)
        # password
        new_pwd = tk.StringVar()
        tk.Label(window_sign_up, text='Please enter password：').place(x=10, y=50)
        tk.Entry(window_sign_up, textvariable=new_pwd, show='*').place(x=150, y=50)
        # reenter password
        new_pwd_confirm = tk.StringVar()
        tk.Label(window_sign_up, text='please reenter password：').place(x=10, y=90)
        tk.Entry(window_sign_up, textvariable=new_pwd_confirm, show='*').place(x=150, y=90)
        bt_confirm_sign_up = tk.Button(window_sign_up, text='register',
                                       command=signtowcg)
        bt_confirm_sign_up.place(x=150, y=130)

    def usr_sign_quit():
        window.destroy()
    bt_login = tk.Button(window_admin, text='login', command=usr_log_in)
    bt_login.place(x=80, y=140)
    bt_logup = tk.Button(window_admin, text='sign up', command=usr_sign_up)
    bt_logup.place(x=150, y=140)
    bt_logquit = tk.Button(window_admin, text='exit', command=usr_sign_quit)
    bt_logquit.place(x=220, y=140)


bt_admin = tk.Button(window,text='Admin',command = admin_sign,bg='#FF9900',height = 2,width =8,relief ='raised' )
bt_admin.place(x=40,y=120)
tk.Label(window,text='Hint:Admin is only used for management').place(x=120,y=130)



# main loop
window.mainloop()